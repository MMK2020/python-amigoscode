from docx import Document
from docx.shared import Inches
import pyttsx3

pyttsx3.speak('Hello')
document = Document()

#Profile picture
document.add_picture('myphoto.jpg',
                     width=Inches(2.0))

#name phone number and email details
name = input('What is your name? ')
pyttsx3.speak('Hello ' + name + 'How are you today')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(name + ' | ' +
 phone_number + ' | ' +email)

# about me
document.add_heading('About Me')
about_me = input ('Tell me about yourself? ')
document.add_paragraph(about_me)

# Work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic

experience_details = input('Describe your experience at ' + company)
p.add_run(experience_details)



# more experiences
while True:
    has_more_experiences = input('Do yu have more experiences? YES or NO ')
    if has_more_experiences.lower() == 'yes':

        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic

        experience_details = input('Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

#Work skills
document.add_heading('Skills')
p = document.add_paragraph(style='List Bullet')

skill_details = input('Enter a work skill ')
p.add_run(skill_details)

#More skills

while True:
    has_more_skills = input('Do you have more skills? YES or NO ')
    if has_more_skills.lower() == 'yes':

        p = document.add_paragraph(style='List Bullet')
        skill_details = input('Enter a work skill ')

        p.add_run(skill_details)
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = " CV generated using Amigoscode and Intuit Quickbooks"


document.save('cv.docx')